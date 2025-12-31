import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Inches
from PIL import Image
import os
import random


template_path = None
logo_path = None


def select_template():
    global template_path
    file_path = filedialog.askopenfilename(
        title="Выберите шаблон документа",
        filetypes=[("Word files", "*.docx")]
    )
    if not file_path:
        return
    template_path = file_path
    template_label.config(text=os.path.basename(file_path))


def select_logo():
    global logo_path
    file_path = filedialog.askopenfilename(
        title="Выберите изображение",
        filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.bmp;*.gif")]
    )
    if not file_path:
        return
    logo_path = file_path
    logo_label.config(text=os.path.basename(file_path))


def resize_logo(src, dst, size=(125, 125)):
    img = Image.open(src)
    img = img.resize(size, Image.Resampling.LANCZOS)
    img.save(dst)


def generate_doc():
    if not template_path:
        messagebox.showwarning("Ошибка", "Сначала выбери шаблон!")
        return

    name = name_input.get()
    age = age_input.get()
    date = date_input.get()
    filename = file_input.get()
    gender = gender_input.get()
    some_key = "".join(random.choice("01") for _ in range(20))

    if not name or not age or not date or not filename:
        messagebox.showwarning("Ошибка", "Заполни все поля!")
        return

    values = {
        "{{NAME}}": name,
        "{{AGE}}": age,
        "{{GENDER}}": gender,
        "{{DATE}}": date,
        "{{KEY}}": some_key,
    }

    doc = Document(template_path)

    for p in doc.paragraphs:
        for key, val in values.items():
            if key in p.text:
                p.text = p.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    if logo_path:
        temp_logo = "resized_logo.png"
        resize_logo(logo_path, temp_logo)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{LOGO}}" in cell.text:
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.alignment = 2
                        run = p.add_run()
                        run.add_picture(temp_logo, width=Inches(1.3), height=Inches(1.3))

        os.remove(temp_logo)

    output_name = f"{filename}.docx"
    doc.save(output_name)

    messagebox.showinfo("Готово", f"Документ сохранён как:\n{output_name}")


root = tk.Tk()
root.title("Генератор документов")
root.geometry("620x420")

main = tk.Frame(root, padx=15, pady=15)
main.pack(fill="both", expand=True)

files_frame = tk.LabelFrame(main, text="Файлы", padx=10, pady=10)
files_frame.pack(fill="x")

tk.Button(files_frame, text="Выбрать шаблон", width=18, command=select_template)\
    .grid(row=0, column=0, padx=5)
template_label = tk.Label(files_frame, text="шаблон не выбран", fg="gray")
template_label.grid(row=0, column=1, sticky="w")

tk.Button(files_frame, text="Выбрать логотип", width=18, command=select_logo)\
    .grid(row=1, column=0, padx=5, pady=4)
logo_label = tk.Label(files_frame, text="логотип не выбран", fg="gray")
logo_label.grid(row=1, column=1, sticky="w")

form = tk.LabelFrame(main, text="Данные", padx=10, pady=10)
form.pack(fill="x", pady=10)

tk.Label(form, text="Дата:").grid(row=0, column=0, sticky="w")
date_input = tk.Entry(form)
date_input.grid(row=0, column=1, sticky="ew", padx=5)

tk.Label(form, text="Имя:").grid(row=1, column=0, sticky="w")
name_input = tk.Entry(form)
name_input.grid(row=1, column=1, sticky="ew", padx=5)

tk.Label(form, text="Пол:").grid(row=2, column=0, sticky="w")
gender_input = tk.Entry(form)
gender_input.grid(row=2, column=1, sticky="ew", padx=5)

tk.Label(form, text="Возраст:").grid(row=3, column=0, sticky="w")
age_input = tk.Entry(form)
age_input.grid(row=3, column=1, sticky="ew", padx=5)

tk.Label(form, text="Название файла (без .docx):").grid(row=4, column=0, sticky="w")
file_input = tk.Entry(form)
file_input.grid(row=4, column=1, sticky="ew", padx=5)

form.columnconfigure(1, weight=1)

tk.Button(main, text="Создать документ", width=22, command=generate_doc)\
    .pack(pady=10)

root.mainloop()
